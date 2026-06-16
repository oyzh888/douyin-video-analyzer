#!/usr/bin/env python3
"""
标准化脚本生成模块
基于分析结果生成符合格式的脚本文档
"""

import json
from pathlib import Path
from datetime import datetime

def generate_script_docx(metadata, analysis, output_path):
    """
    生成标准化脚本文档（.docx格式）
    
    参数:
        metadata: 视频元数据
        analysis: Gemini分析结果
        output_path: 输出文件路径
    
    返回: 生成的脚本内容
    """
    try:
        import docx
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("⚠️  python-docx模块未安装，生成文本格式")
        return generate_script_text(metadata, analysis, output_path.replace('.docx', '.txt'))
    
    print(f"📝 生成标准化脚本: {output_path}")
    
    # 创建文档
    doc = Document()
    
    # 设置基本样式
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(12)
    
    # 1. 标题部分
    title = analysis.get('title', metadata.get('title', '短视频脚本'))
    doc.add_heading(f"剧名：《{title}》", level=0)
    
    # 整体剧情概括
    summary = analysis.get('video_summary', '')
    if summary:
        doc.add_heading("整体剧情概括", level=1)
        doc.add_paragraph(summary)
    
    # 2. 分集内容
    doc.add_heading("第一集", level=1)
    
    # 脚本表
    script_table = analysis.get('script_table', [])
    if script_table:
        for i, scene in enumerate(script_table):
            scene_num = i + 1
            time_range = scene.get('time_range', f'00:00-00:00')
            visual = scene.get('visual_content_inferred', '')
            actions = scene.get('actions_inferred', '')
            dialogue = scene.get('dialogue', '')
            
            # 场景标题
            doc.add_heading(f"{scene_num}-{scene_num} 日 内 场景", level=2)
            
            # 画面描述（△符号）
            if visual:
                para = doc.add_paragraph()
                para.add_run(f"△ {visual}")
            
            if actions:
                para = doc.add_paragraph()
                para.add_run(f"△ {actions}")
            
            # 对话
            if dialogue:
                para = doc.add_paragraph()
                para.add_run(dialogue)
            
            doc.add_paragraph()  # 空行分隔
    
    # 3. 角色分析
    doc.add_heading("【角色分析】", level=1)
    
    core_points = analysis.get('core_points', [])
    if core_points:
        for point in core_points:
            doc.add_paragraph(f"• {point}", style='ListBullet')
    
    # 4. 教育意义/文化价值
    doc.add_heading("【教育意义/文化价值】", level=1)
    
    educational_value = analysis.get('educational_value', '')
    if educational_value:
        doc.add_paragraph(educational_value)
    
    cultural_value = analysis.get('cultural_value', '')
    if cultural_value:
        doc.add_paragraph(cultural_value)
    
    # 5. 技术参数
    doc.add_heading("【技术参数】", level=1)
    
    params = [
        f"时长: {metadata.get('duration', '未知')}秒",
        f"平台: {metadata.get('platform', '抖音')}",
        f"作者: {metadata.get('author', '未知')}",
        f"方言特色: {analysis.get('dialect_features', '普通话')}",
        f"情感节奏: {analysis.get('emotional_value', '温情叙事')}",
        f"目标受众: {analysis.get('target_audience', '大众')}",
        f"改编潜力: {analysis.get('adaptation_potential', '中等')}"
    ]
    
    for param in params:
        doc.add_paragraph(param)
    
    # 6. 可替换部分
    replaceable_parts = analysis.get('replaceable_parts', [])
    if replaceable_parts:
        doc.add_heading("【可替换部分】", level=1)
        for part in replaceable_parts:
            doc.add_paragraph(f"• {part}", style='ListBullet')
    
    # 7. 生成信息
    doc.add_paragraph()
    doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"分析工具: 抖音视频分析技能 v1.0")
    doc.add_paragraph(f"源视频: {metadata.get('url', '未知')}")
    
    # 保存文档
    doc.save(output_path)
    
    print(f"✅ 脚本生成完成: {output_path}")
    
    # 返回文本内容供预览
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    return "\n".join(full_text)

def generate_script_text(metadata, analysis, output_path):
    """
    生成文本格式脚本（备选方案）
    """
    print(f"📝 生成文本格式脚本: {output_path}")
    
    lines = []
    
    # 1. 标题部分
    title = analysis.get('title', metadata.get('title', '短视频脚本'))
    lines.append(f"剧名：《{title}》")
    lines.append("")
    
    # 整体剧情概括
    summary = analysis.get('video_summary', '')
    if summary:
        lines.append("整体剧情概括：")
        lines.append(summary)
        lines.append("")
    
    # 2. 分集内容
    lines.append("第一集：")
    lines.append("")
    
    # 脚本表
    script_table = analysis.get('script_table', [])
    if script_table:
        for i, scene in enumerate(script_table):
            scene_num = i + 1
            time_range = scene.get('time_range', f'00:00-00:00')
            visual = scene.get('visual_content_inferred', '')
            actions = scene.get('actions_inferred', '')
            dialogue = scene.get('dialogue', '')
            
            lines.append(f"{scene_num}-{scene_num} 日 内 场景")
            
            if visual:
                lines.append(f"△ {visual}")
            
            if actions:
                lines.append(f"△ {actions}")
            
            if dialogue:
                lines.append(dialogue)
            
            lines.append("")
    
    # 3. 角色分析
    lines.append("【角色分析】")
    core_points = analysis.get('core_points', [])
    if core_points:
        for point in core_points:
            lines.append(f"• {point}")
    lines.append("")
    
    # 4. 教育意义/文化价值
    lines.append("【教育意义/文化价值】")
    educational_value = analysis.get('educational_value', '')
    if educational_value:
        lines.append(educational_value)
    
    cultural_value = analysis.get('cultural_value', '')
    if cultural_value:
        lines.append(cultural_value)
    lines.append("")
    
    # 5. 技术参数
    lines.append("【技术参数】")
    params = [
        f"时长: {metadata.get('duration', '未知')}秒",
        f"平台: {metadata.get('platform', '抖音')}",
        f"作者: {metadata.get('author', '未知')}",
        f"方言特色: {analysis.get('dialect_features', '普通话')}",
        f"情感节奏: {analysis.get('emotional_value', '温情叙事')}",
        f"目标受众: {analysis.get('target_audience', '大众')}",
        f"改编潜力: {analysis.get('adaptation_potential', '中等')}"
    ]
    
    for param in params:
        lines.append(param)
    lines.append("")
    
    # 6. 可替换部分
    replaceable_parts = analysis.get('replaceable_parts', [])
    if replaceable_parts:
        lines.append("【可替换部分】")
        for part in replaceable_parts:
            lines.append(f"• {part}")
        lines.append("")
    
    # 7. 生成信息
    lines.append(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"分析工具: 抖音视频分析技能 v1.0")
    lines.append(f"源视频: {metadata.get('url', '未知')}")
    
    # 保存文件
    content = "\n".join(lines)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(content)
    
    print(f"✅ 文本脚本生成完成: {output_path}")
    
    return content

def generate_html_report(metadata, analysis, output_path):
    """
    生成HTML分析报告
    """
    print(f"📊 生成HTML报告: {output_path}")
    
    title = analysis.get('title', metadata.get('title', '视频分析报告'))
    
    html_template = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title} - 分析报告</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
            margin: 0;
            padding: 20px;
            background: #f5f5f5;
        }}
        .container {{
            max-width: EI
px;
            margin: 0 auto;
            background: white;
            padding: 30px;
            border-radius: 10px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #333;
            border-bottom: 2px solid #4CAF50;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #555;
            margin-top: 20px;
        }}
        .card {{
            background: #f9f9f9;
            padding: 15px;
            margin: 15px 0;
            border-left: 4px solid #4CAF50;
            border-radius: 0 5px 5px 0;
        }}
        .metadata {{
            color: #666;
            font-size: 0.9em;
            background: #f0f7ff;
            padding: 15px;
            border-radius: 8px;
        }}
        .scene {{
            border: 1px solid #ddd;
            padding: 15px;
            margin: 10px 0;
            border-radius: 5px;
        }}
        .time {{
            font-weight: bold;
            color: #2196F3;
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>{title} - 分析报告</h1>
        
        <div class="metadata">
            <p><strong>视频标题：</strong>{metadata.get('title', '未知')}</p>
            <p><strong>作者：</strong>{metadata.get('author', '未知')}</p>
            <p><strong>时长：</strong>{metadata.get('duration', '未知')}秒</p>
            <p><strong>平台：</strong>{metadata.get('platform', '抖音')}</p>
            <p><strong>分析时间：</strong>{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        </div>
        
        <h2>📝 视频整体内容总结</h2>
        <div class="card">{analysis.get('video_summary', '')}</div>
        
        <h2>💥 核心爆点</h2>
        <div class="card">
            <ul>
                {"".join([f'<li>{point}</li>' for point in analysis.get('core_points', [])])}
            </ul>
        </div>
        
        <h2>📖 脚本结构</h2>
        {generate_html_script_table(analysis.get('script_table', []))}
        
        <h2>🎯 深度分析</h2>
        <div class="card">
            <p><strong>视频类型：</strong>{analysis.get('genre_guess', '未知')}</p>
            <p><strong>目标受众：</strong>{analysis.get('target_audience', '未知')}</p>
            <p><strong>教育意义：</strong>{analysis.get('educational_value', '')}</p>
            <p><strong>文化价值：</strong>{analysis.get('cultural_value', '')}</p>
            <p><strong>方言特色：</strong>{analysis.get('dialect_features', '')}</p>
        </div>
        
        <div style="margin-top: 30px; padding-top: 20px; border-top: 1px solid #eee; color: #666; font-size: 0.9em;">
            <p>生成工具：抖音视频分析技能 v1.0</p>
            <p>源视频：{metadata.get('url', '未知')}</p>
        </div>
    </div>
</body>
</html>"""
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html_template)
    
    print(f"✅ HTML报告生成完成: {output_path}")
    
    return output_path

def generate_html_script_table(script_table):
    """生成脚本表的HTML"""
    if not script_table:
        return '<div class="card">暂无详细的脚本表数据</div>'
    
    html = '<div class="script-table">'
    for scene in script_table:
        html += f'''
        <div class="scene">
            <div class="time">时间：{scene.get('time_range', '00:00-00:00')}</div>
            <p><strong>画面内容：</strong>{scene.get('visual_content_inferred', '')}</p>
            <p><strong>动作：</strong>{scene.get('actions_inferred', '')}</p>
            <p><strong>对白：</strong>{scene.get('dialogue', '')}</p>
        </div>
        '''
    html += '</div>'
    
    return html

if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 3:
        print("用法: python generate_script.py <元数据JSON文件> <分析结果JSON文件> [输出文件]")
        print("示例: python generate_script.py metadata.json analysis.json script.docx")
        sys.exit(1)
    
    metadata_file = sys.argv[1]
    analysis_file = sys.argv[2]
    output_file = sys.argv[3] if len(sys.argv) > 3 else "output.docx"
    
    try:
        with open(metadata_file, 'r', encoding='utf-8') as f:
            metadata = json.load(f)
        
        with open(analysis_file, 'r', encoding='utf-8') as f:
            analysis = json.load(f)
        
        result = generate_script_docx(metadata, analysis, output_file)
        print(f"\n✅ 脚本生成成功!")
        print(f"📄 输出文件: {output_file}")
        
        # 生成HTML报告
        html_path = output_file.replace('.docx', '.html').replace('.txt', '.html')
        generate_html_report(metadata, analysis, html_path)
        
    except Exception as e:
        print(f"❌ 脚本生成失败: {e}")
        import traceback
        traceback.print_exc()